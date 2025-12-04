import io
import base64
from datetime import datetime
import numpy as np

# Matplotlib (Object-Oriented API for thread safety)
from matplotlib.figure import Figure
from matplotlib.backends.backend_agg import FigureCanvasAgg as FigureCanvas
import matplotlib.pyplot as plt

# Para PDF
from xhtml2pdf import pisa

# Para Word
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def crear_radar_matplotlib(categorias, valores_evaluado, valores_empresa=None, nombre_evaluado='Evaluado'):
    """Genera un gráfico de radar usando Matplotlib."""
    # Cerrar el ciclo
    N = len(categorias)
    angles = [n / float(N) * 2 * np.pi for n in range(N)]
    angles += angles[:1]
    
    valores_evaluado = list(valores_evaluado) + [valores_evaluado[0]]
    if valores_empresa:
        valores_empresa = list(valores_empresa) + [valores_empresa[0]]
    
    # Crear figura
    fig = Figure(figsize=(8, 8), dpi=100)
    canvas = FigureCanvas(fig)
    ax = fig.add_subplot(111, polar=True)
    
    # Configurar ejes
    ax.set_theta_offset(np.pi / 2)
    ax.set_theta_direction(-1)
    
    # Etiquetas
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categorias, size=10)
    
    # Eje Y
    ax.set_rlabel_position(0)
    plt.yticks([1, 2, 3, 4, 5], ["1", "2", "3", "4", "5"], color="grey", size=7)
    ax.set_ylim(0, 5)
    
    # Plot Evaluado
    ax.plot(angles, valores_evaluado, linewidth=2, linestyle='solid', label=nombre_evaluado, color='#667eea')
    ax.fill(angles, valores_evaluado, '#667eea', alpha=0.25)
    
    # Plot Empresa (si existe)
    if valores_empresa:
        ax.plot(angles, valores_empresa, linewidth=1, linestyle='dashed', label='Promedio Empresa', color='#6c757d')
    
    # Plot Niveles Referencia
    valores_sobresaliente = [4.5] * (N + 1)
    ax.plot(angles, valores_sobresaliente, linewidth=1, linestyle='dotted', color='#28a745', alpha=0.5)
    
    valores_aceptable = [3.5] * (N + 1)
    ax.plot(angles, valores_aceptable, linewidth=1, linestyle='dotted', color='#ffc107', alpha=0.5)

    ax.legend(loc='upper right', bbox_to_anchor=(1.1, 1.1))
    
    # Guardar
    buf = io.BytesIO()
    fig.savefig(buf, format='png', bbox_inches='tight')
    buf.seek(0)
    return buf.getvalue()

def crear_dona_matplotlib(valor, color, titulo):
    """Genera un gráfico de dona usando Matplotlib."""
    fig = Figure(figsize=(4, 4), dpi=100)
    canvas = FigureCanvas(fig)
    ax = fig.add_subplot(111)
    
    sizes = [valor, 5 - valor]
    colors = [color, '#f0f0f0']
    
    # Dona
    wedges, texts = ax.pie(sizes, colors=colors, startangle=90, wedgeprops=dict(width=0.3))
    
    # Texto central
    ax.text(0, 0, f"{valor:.2f}", ha='center', va='center', fontsize=20, fontweight='bold', color='#333')
    
    # Título (opcional, mejor manejarlo fuera)
    # ax.set_title(titulo)
    
    buf = io.BytesIO()
    fig.savefig(buf, format='png', bbox_inches='tight', transparent=True)
    buf.seek(0)
    return buf.getvalue()

def crear_perfil_competencias_matplotlib(categorias, valores_evaluado):
    """Genera un gráfico de radar (Perfil de Competencias) comparando con el estándar."""
    N = len(categorias)
    angles = [n / float(N) * 2 * np.pi for n in range(N)]
    angles += angles[:1]
    
    valores_evaluado = list(valores_evaluado) + [valores_evaluado[0]]
    valores_estandar = [3.5] * (N + 1)
    
    fig = Figure(figsize=(5, 5), dpi=100)
    canvas = FigureCanvas(fig)
    ax = fig.add_subplot(111, polar=True)
    
    ax.set_theta_offset(np.pi / 2)
    ax.set_theta_direction(-1)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categorias, size=9)
    ax.set_rlabel_position(0)
    plt.yticks([1, 2, 3, 4, 5], ["1", "2", "3", "4", "5"], color="grey", size=7)
    ax.set_ylim(0, 5)
    
    # Evaluado
    ax.plot(angles, valores_evaluado, linewidth=2, linestyle='solid', label='Evaluación Actual', color='#667eea')
    ax.fill(angles, valores_evaluado, '#667eea', alpha=0.25)
    
    # Estándar
    ax.plot(angles, valores_estandar, linewidth=2, linestyle='dashed', label='Estándar Mínimo (3.5)', color='#ff6b6b')
    
    ax.legend(loc='upper right', bbox_to_anchor=(1.3, 1.1), fontsize='small')
    
    buf = io.BytesIO()
    fig.savefig(buf, format='png', bbox_inches='tight')
    buf.seek(0)
    return buf.getvalue()

def crear_madurez_habilidades_matplotlib(categorias, valores_evaluado, valores_empresa=None, nombre_evaluado='Evaluado'):
    """Genera un gráfico de radar para Análisis de Madurez con benchmarks específicos."""
    N = len(categorias)
    angles = [n / float(N) * 2 * np.pi for n in range(N)]
    angles += angles[:1]
    
    valores_evaluado = list(valores_evaluado) + [valores_evaluado[0]]
    if valores_empresa:
        valores_empresa = list(valores_empresa) + [valores_empresa[0]]
    
    # Benchmarks
    val_sobresaliente = [4.5] * (N + 1)
    val_aceptable = [3.5] * (N + 1)

    fig = Figure(figsize=(6.5, 6.5), dpi=100)
    canvas = FigureCanvas(fig)
    ax = fig.add_subplot(111, polar=True)
    
    ax.set_theta_offset(np.pi / 2)
    ax.set_theta_direction(-1)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categorias, size=10)
    ax.set_rlabel_position(0)
    plt.yticks([1, 2, 3, 4, 5], ["1", "2", "3", "4", "5"], color="grey", size=7)
    ax.set_ylim(0, 5)
    
    # Plot Benchmarks
    ax.plot(angles, val_sobresaliente, linewidth=1, linestyle='--', label='Sobresaliente (4.5)', color='#28a745')
    ax.plot(angles, val_aceptable, linewidth=1, linestyle='--', label='Aceptable (3.5)', color='#ffc107')

    # Plot Empresa
    if valores_empresa:
        ax.plot(angles, valores_empresa, linewidth=1.5, linestyle='-.', label='Promedio Empresa', color='#6c757d')

    # Plot Evaluado
    ax.plot(angles, valores_evaluado, linewidth=2, linestyle='solid', label=nombre_evaluado, color='#667eea')
    ax.fill(angles, valores_evaluado, '#667eea', alpha=0.25)
    
    ax.legend(loc='upper right', bbox_to_anchor=(1.3, 1.1), fontsize='small')
    
    buf = io.BytesIO()
    fig.savefig(buf, format='png', bbox_inches='tight')
    buf.seek(0)
    return buf.getvalue()

def crear_radar_general_matplotlib(categorias, valores_evaluado):
    """Genera el radar general (Evaluado vs Estándar)."""
    N = len(categorias)
    angles = [n / float(N) * 2 * np.pi for n in range(N)]
    angles += angles[:1]
    
    valores_evaluado = list(valores_evaluado) + [valores_evaluado[0]]
    valores_estandar = [3.5] * (N + 1)
    
    fig = Figure(figsize=(6, 6), dpi=100)
    canvas = FigureCanvas(fig)
    ax = fig.add_subplot(111, polar=True)
    
    ax.set_theta_offset(np.pi / 2)
    ax.set_theta_direction(-1)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categorias, size=9)
    ax.set_rlabel_position(0)
    plt.yticks([1, 2, 3, 4, 5], ["1", "2", "3", "4", "5"], color="grey", size=7)
    ax.set_ylim(0, 5)
    
    # Evaluado
    ax.plot(angles, valores_evaluado, linewidth=2, linestyle='solid', label='Evaluación Actual', color='#667eea')
    ax.fill(angles, valores_evaluado, '#667eea', alpha=0.25)
    
    # Estándar
    ax.plot(angles, valores_estandar, linewidth=2, linestyle='dashed', label='Estándar Mínimo (3.5)', color='#ff6b6b')
    
    ax.legend(loc='upper right', bbox_to_anchor=(1.3, 1.1), fontsize='small')
    
    buf = io.BytesIO()
    fig.savefig(buf, format='png', bbox_inches='tight')
    buf.seek(0)
    return buf.getvalue()

def crear_matriz_9box_matplotlib(potencial, desempeno, evaluado, cuadrante_info):
    """Genera la matriz 9-Box."""
    fig = Figure(figsize=(6, 6), dpi=100)
    canvas = FigureCanvas(fig)
    ax = fig.add_subplot(111)
    
    # Ejes y límites
    ax.set_xlim(0, 5)
    ax.set_ylim(0, 5)
    ax.set_xlabel("Potencial")
    ax.set_ylabel("Desempeño")
    
    # Líneas divisorias
    ax.axhline(y=4.0, color='gray', linestyle='--', alpha=0.5)
    ax.axvline(x=4.0, color='gray', linestyle='--', alpha=0.5)
    
    # Punto del evaluado
    color = cuadrante_info.get('color_cuadrante', '#667eea')
    ax.scatter([potencial], [desempeno], s=200, c=color, edgecolors='white', linewidths=2, zorder=5)
    ax.text(potencial, desempeno + 0.15, evaluado[:15], ha='center', fontsize=9, fontweight='bold')
    
    # Fondo (opcional, para dar contexto)
    ax.grid(True, linestyle=':', alpha=0.3)
    
    buf = io.BytesIO()
    fig.savefig(buf, format='png', bbox_inches='tight')
    buf.seek(0)
    return buf.getvalue()

def crear_comparacion_barras_matplotlib(categorias, grupos_data, colores_grupos):
    """Genera gráfico de barras agrupadas para comparación."""
    fig = Figure(figsize=(10, 5), dpi=100)
    canvas = FigureCanvas(fig)
    ax = fig.add_subplot(111)
    
    x = np.arange(len(categorias))
    width = 0.8 / len(grupos_data)
    
    for i, (grupo, valores) in enumerate(grupos_data.items()):
        offset = (i - len(grupos_data)/2) * width + width/2
        color = colores_grupos.get(grupo, '#6c757d')
        ax.bar(x + offset, valores, width, label=grupo, color=color)
        
    ax.set_xticks(x)
    ax.set_xticklabels(categorias, rotation=15, ha='right', fontsize=9)
    ax.set_ylim(0, 5)
    ax.legend(loc='upper center', bbox_to_anchor=(0.5, 1.15), ncol=len(grupos_data), fontsize='small')
    ax.grid(axis='y', linestyle='--', alpha=0.3)
    
    buf = io.BytesIO()
    fig.savefig(buf, format='png', bbox_inches='tight')
    buf.seek(0)
    return buf.getvalue()

def generar_imagenes_matplotlib(datos):
    """Genera todas las imágenes necesarias usando Matplotlib."""
    imagenes = {}
    raw = datos['data_raw']
    meta = datos['meta']
    kpis = datos['kpis']
    textos = datos['textos']
    
    cats = list(raw['promedios_categorias'].keys())
    vals = list(raw['promedios_categorias'].values())
    
    # 1. Radar General (REMOVED)
    # 2. Radar Avanzado (REMOVED)
    empresa = raw.get('promedios_empresa_cat', [])
    
    # 3. Matriz 9-Box
    # Calcular potencial (promedio de ciertas competencias)
    cats_potencial = ['Liderazgo', 'Innovación y Creatividad', 'Toma de Decisiones']
    vals_potencial = [raw['promedios_categorias'].get(c, 0) for c in cats_potencial if c in raw['promedios_categorias']]
    potencial = sum(vals_potencial) / len(vals_potencial) if vals_potencial else 0
    desempeno = kpis['calificacion_final']
    
    print("Generando Matriz 9-Box (Matplotlib)...")
    img_matriz = crear_matriz_9box_matplotlib(potencial, desempeno, meta['evaluado'], textos)
    imagenes['matriz'] = img_matriz

    # 4. Comparación (Barras)
    # Necesitamos reconstruir los datos de grupos si no están en raw.
    # En app.py 'grupos' no se pasa directamente en data_raw, pero podemos intentar inferirlo o pasarlo.
    # Por ahora, si no está disponible, omitimos o usamos dummy si es crítico.
    # REVISIÓN: app.py NO pasa 'grupos' en data_raw. Necesitamos modificar app.py para pasarlo o calcularlo aquí.
    # Sin embargo, modificar app.py es riesgoso si no tenemos el DF.
    # Vemos que 'data_raw' tiene 'final_por_comp', 'colores_categorias', 'promedios_categorias'.
    # Faltan los promedios por grupo.
    # Voy a asumir que se agregará a data_raw en app.py o lo omitimos por ahora.
    # UPDATE: Voy a agregar 'promedios_por_grupo' a data_raw en app.py en el siguiente paso.
    if 'promedios_por_grupo' in raw:
        print("Generando Comparación (Matplotlib)...")
        img_comp = crear_comparacion_barras_matplotlib(cats, raw['promedios_por_grupo'], {
            'Autoevaluación': '#17a2b8', 'Jefe Inmediato': '#dc3545', 'Colegas': '#ffc107', 'Subordinados': '#28a745'
        })
        imagenes['comparacion'] = img_comp
    
    # 5. Perfil de Competencias
    if cats:
        print("Generando Perfil de Competencias (Matplotlib)...")
        img_perfil = crear_perfil_competencias_matplotlib(cats, vals)
        imagenes['perfil_competencias'] = img_perfil

    # 6. Análisis de Madurez
    if cats:
        print("Generando Análisis de Madurez (Matplotlib)...")
        # Usamos la misma data de empresa si existe
        img_madurez = crear_madurez_habilidades_matplotlib(cats, vals, empresa if empresa else None, meta.get('evaluado', 'Evaluado'))
        imagenes['madurez'] = img_madurez
    
    # 7. Donas por Categoría
    colores = raw['colores_categorias']
    for cat, val in raw['promedios_categorias'].items():
        # print(f"Generando Dona {cat} (Matplotlib)...") # Reduce noise
        color = colores.get(cat, '#667eea')
        img_dona = crear_dona_matplotlib(val, color, cat)
        imagenes[f'cat_{cat}'] = img_dona
        
    return imagenes

def generar_pdf(datos):
    """Genera PDF usando imágenes de Matplotlib."""
    # Generar imágenes
    imagenes_bytes = generar_imagenes_matplotlib(datos)
    
    # Convertir a base64
    imagenes_b64 = {}
    for k, v in imagenes_bytes.items():
        if v:
            encoded = base64.b64encode(v).decode('utf-8')
            imagenes_b64[k] = f"data:image/png;base64,{encoded}"
        else:
            imagenes_b64[k] = ""

    # Datos
    meta = datos['meta']
    kpis = datos['kpis']
    textos = datos['textos']
    
    # Grid de categorías
    categorias_html = ""
    cats_keys = [k for k in imagenes_b64.keys() if k.startswith('cat_')]
    
    for i in range(0, len(cats_keys), 2):
        row_keys = cats_keys[i:i+2]
        categorias_html += "<tr>"
        for key in row_keys:
            cat_name = key.replace('cat_', '')
            img_src = imagenes_b64.get(key, '')
            categorias_html += f"""
            <td style="text-align: center; border: none; padding: 10px;">
                <h4 style="color: #667eea; margin-bottom: 5px;">{cat_name}</h4>
                <p style="font-size: 10px; color: #666; margin-bottom: 10px;">Nivel de dominio alcanzado en <strong>{cat_name}</strong>.</p>
                <img src="{img_src}" style="width: 200px; height: auto;" />
            </td>
            """
        categorias_html += "</tr>"

    html_content = f"""
    <html>
    <head>
        <style>
            @page {{
                size: letter;
                margin: 2.5cm;
                @frame footer_frame {{
                    -pdf-frame-content: footerContent;
                    bottom: 1cm;
                    margin-left: 2.5cm;
                    margin-right: 2.5cm;
                    height: 1cm;
                }}
            }}
            body {{ font-family: Helvetica, sans-serif; color: #333; line-height: 1.5; }}
            h1 {{ color: #2c3e50; border-bottom: 3px solid #667eea; padding-bottom: 15px; font-size: 28px; margin-bottom: 25px; text-transform: uppercase; letter-spacing: 1px; }}
            h2 {{ color: #667eea; margin-top: 35px; font-size: 20px; border-bottom: 2px solid #f0f0f0; padding-bottom: 8px; font-weight: bold; }}
            h3 {{ color: #17a2b8; font-size: 18px; margin-top: 0; font-weight: bold; }}
            h4 {{ color: #6c757d; font-size: 14px; margin-bottom: 5px; font-weight: bold; }}
            .header {{ text-align: center; margin-bottom: 50px; background-color: #f8f9fa; padding: 20px; border-radius: 10px; }}
            .kpi-container {{ text-align: center; margin-bottom: 40px; }}
            .kpi-box {{ 
                background-color: #fff; padding: 15px; border: 1px solid #e9ecef; border-radius: 8px; box-shadow: 2px 2px 5px rgba(0,0,0,0.05);
                text-align: center; display: inline-block; width: 20%; margin: 1%; vertical-align: top;
            }}
            .score {{ font-size: 24px; font-weight: bold; color: #2c3e50; margin-top: 5px; }}
            .label {{ font-size: 11px; color: #6c757d; text-transform: uppercase; letter-spacing: 1px; font-weight: bold; }}
            .aptitud-box {{
                background-color: #eef2f7; border-left: 6px solid {textos.get('color_aptitud', '#667eea')};
                padding: 25px; margin: 30px 0; border-radius: 0 8px 8px 0; box-shadow: 2px 2px 5px rgba(0,0,0,0.05);
            }}
            .img-container {{ text-align: center; margin: 35px 0; page-break-inside: avoid; }}
            img {{ max-width: 100%; height: auto; }}
            .footer {{ text-align: center; font-size: 10px; color: #aaa; border-top: 1px solid #eee; padding-top: 15px; margin-top: 30px; }}
            table {{ width: 100%; border-collapse: collapse; margin-top: 20px; font-size: 12px; }}
            th, td {{ border: 1px solid #dee2e6; padding: 10px; text-align: left; vertical-align: top; }}
            th {{ background-color: #f1f3f5; color: #495057; font-weight: bold; text-transform: uppercase; font-size: 11px; }}
            tr:nth-child(even) {{ background-color: #f8f9fa; }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1 style="border: none; margin-bottom: 10px;">Reporte de Evaluación 360°</h1>
            <p style="font-size: 18px; color: #2c3e50;"><strong>{meta['evaluado']}</strong></p>
            <p>Fecha: {datetime.now().strftime('%d/%m/%Y')}</p>
        </div>

        <div class="kpi-container">
            <div class="kpi-box"><div class="label">Calificación</div><div class="score">{kpis['calificacion_final']:.2f}</div></div>
            <div class="kpi-box"><div class="label">Cumplimiento</div><div class="score">{kpis['nivel_cumplimiento']:.0f}%</div></div>
            <div class="kpi-box"><div class="label">Consistencia</div><div class="score">{kpis['consistencia']:.1f}</div></div>
            <div class="kpi-box"><div class="label">Percentil</div><div class="score">{100 - kpis['percentil']:.0f}%</div></div>
        </div>

        <div class="aptitud-box">
            <h3>{textos['estado_aptitud']}</h3>
            <p>{textos['mensaje_aptitud']}</p>
            <p><strong>Recomendación RH:</strong> {textos['recomendacion_rh']}</p>
        </div>

        <h2>Análisis de Competencias</h2>
        <!-- Radares removidos a petición del usuario -->

        <h2>Perfil de Competencias</h2>
        <p style="text-align: justify; font-size: 12px; color: #666; margin-bottom: 20px;">
            Este gráfico muestra el desempeño del evaluado en cada competencia (área azul) comparado con el estándar mínimo requerido (línea roja discontinua). Permite visualizar rápidamente las fortalezas y áreas que requieren atención.
        </p>
        <div class="img-container">
            <img src="{imagenes_b64.get('perfil_competencias', '')}" style="width: 400px;" />
        </div>

        <h2>Análisis de Madurez por Habilidades</h2>
        <p style="text-align: justify; font-size: 12px; color: #666; margin-bottom: 20px;">
            Este gráfico compara el desempeño de <strong>{meta['evaluado']}</strong> (área azul) frente al promedio de la empresa (línea gris discontinua) y los niveles de referencia: Sobresaliente (4.5, línea verde) y Aceptable (3.5, línea amarilla). Permite identificar rápidamente las brechas de madurez en cada competencia.
        </p>
        <div class="img-container">
            <img src="{imagenes_b64.get('madurez', '')}" style="width: 400px;" />
        </div>
        
        <pdf:nextpage />
        <h2>Matriz de Talento (9-Box)</h2>
        <p style="text-align: justify; font-size: 12px; color: #666; margin-bottom: 20px;">
            Esta matriz ubica al evaluado en función de su <strong>Desempeño</strong> (Eje Y) y <strong>Potencial</strong> (Eje X). Permite identificar si el colaborador es un talento clave, requiere desarrollo o está en una posición adecuada a sus capacidades actuales.
        </p>
        <div class="img-container">
            <img src="{imagenes_b64.get('matriz', '')}" style="width: 450px;" />
        </div>

        <h2>Comparativa por Grupo</h2>
        <br/><br/><br/><br/><br/><br/>
        <p style="text-align: justify; font-size: 12px; color: #666; margin-bottom: 20px;">
            Este gráfico de barras permite contrastar la autoevaluación con la percepción de otros grupos (Jefe, Colegas, Subordinados), facilitando la identificación de puntos ciegos y áreas de consenso.
        </p>
        <div class="img-container">
            <img src="{imagenes_b64.get('comparacion', '')}" style="width: 600px;" />
        </div>

        <pdf:nextpage />
        <h2>Detalle por Competencia</h2>
        <table style="border: none;">
            {categorias_html}
        </table>

        <pdf:nextpage />
        <h2>Detalle por Competencia</h2>
        <table style="border: none;">
            {categorias_html}
        </table>

        <div id="footerContent" class="footer">
            Página <pdf:pagenumber /> | Generado por Dashboard 360°
        </div>
    </body>
    </html>
    """

    buffer = io.BytesIO()
    pisa_status = pisa.CreatePDF(html_content, dest=buffer)
    if pisa_status.err:
        print(f"Error PDF: {pisa_status.err}")
        return None
    buffer.seek(0)
    return buffer

def generar_word(datos):
    """Genera Word usando imágenes de Matplotlib."""
    imagenes_bytes = generar_imagenes_matplotlib(datos)

    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Título
    titulo = document.add_heading('Reporte de Evaluación 360°', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"Evaluado: {datos['meta']['evaluado']}")
    document.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}")

    # KPIs
    table = document.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    headers = ['Calificación', 'Cumplimiento', 'Consistencia', 'Percentil']
    vals = [
        f"{datos['kpis']['calificacion_final']:.2f}",
        f"{datos['kpis']['nivel_cumplimiento']:.0f}%",
        f"{datos['kpis']['consistencia']:.1f}",
        f"Top {100 - datos['kpis']['percentil']:.0f}%"
    ]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for i, v in enumerate(vals):
        table.rows[1].cells[i].text = v

    document.add_paragraph()

    # Diagnóstico
    document.add_heading('Diagnóstico General', level=1)
    p = document.add_paragraph()
    p.add_run(f"{datos['textos']['estado_aptitud']}\n").bold = True
    p.add_run(f"{datos['textos']['mensaje_aptitud']}\n")
    p.add_run(f"Recomendación RH: {datos['textos']['recomendacion_rh']}")

    # Radar (REMOVED)
    # document.add_heading('Análisis de Competencias', level=1)
    
    # if imagenes_bytes.get('radar_general'):
    #     document.add_heading('Radar General', level=2)
    #     img_stream = io.BytesIO(imagenes_bytes['radar_general'])
    #     document.add_picture(img_stream, width=Inches(4.5))
        
    # if imagenes_bytes.get('radar_avanzado'):
    #     document.add_heading('Radar Avanzado', level=2)
    #     img_stream = io.BytesIO(imagenes_bytes['radar_avanzado'])
    #     document.add_picture(img_stream, width=Inches(6.0))

    # Perfil y Madurez
    # document.add_page_break() # Removed to optimize space
    document.add_heading('Perfil de Competencias', level=1)
    p = document.add_paragraph()
    p.add_run("Este gráfico muestra el desempeño del evaluado en cada competencia (área azul) comparado con el estándar mínimo requerido (línea roja discontinua). Permite visualizar rápidamente las fortalezas y áreas que requieren atención.")
    
    if imagenes_bytes.get('perfil_competencias'):
        img_stream = io.BytesIO(imagenes_bytes['perfil_competencias'])
        document.add_picture(img_stream, width=Inches(4.8))

    document.add_heading('Análisis de Madurez por Habilidades', level=1)
    p = document.add_paragraph()
    p.add_run(f"Este gráfico compara el desempeño de {datos['meta']['evaluado']} frente al promedio de la empresa y los niveles de referencia (Sobresaliente 4.5 y Aceptable 3.5). Permite visualizar la madurez profesional en cada competencia clave.")
    
    if imagenes_bytes.get('madurez'):
        img_stream = io.BytesIO(imagenes_bytes['madurez'])
        document.add_picture(img_stream, width=Inches(4.8))

    # Matriz
    document.add_page_break()
    document.add_heading('Matriz de Talento (9-Box)', level=1)
    p = document.add_paragraph()
    p.add_run("Esta matriz ubica al evaluado en función de su Desempeño (Eje Y) y Potencial (Eje X). Permite identificar si el colaborador es un talento clave, requiere desarrollo o está en una posición adecuada a sus capacidades actuales.")
    if imagenes_bytes.get('matriz'):
        img_stream = io.BytesIO(imagenes_bytes['matriz'])
        document.add_picture(img_stream, width=Inches(5.0))

    # Comparación
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    document.add_heading('Comparativa por Grupo', level=1)
    p = document.add_paragraph()
    p.add_run("Este gráfico de barras permite contrastar la autoevaluación con la percepción de otros grupos (Jefe, Colegas, Subordinados), facilitando la identificación de puntos ciegos y áreas de consenso.")
    if imagenes_bytes.get('comparacion'):
        img_stream = io.BytesIO(imagenes_bytes['comparacion'])
        document.add_picture(img_stream, width=Inches(6.5))

    document.add_page_break()
    document.add_heading('Detalle por Competencia', level=1)
    
    cats_keys = [k for k in imagenes_bytes.keys() if k.startswith('cat_')]
    
    table_cats = document.add_table(rows=(len(cats_keys) + 1) // 2, cols=2)
    table_cats.autofit = False
    
    for i, key in enumerate(cats_keys):
        row = i // 2
        col = i % 2
        cell = table_cats.cell(row, col)
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(key.replace('cat_', '')).bold = True
        p.add_run(f"\nNivel de dominio alcanzado en {key.replace('cat_', '')}.").font.size = Pt(9)
        
        if imagenes_bytes.get(key):
            img_stream = io.BytesIO(imagenes_bytes[key])
            run = p.add_run()
            run.add_break()
            run.add_picture(img_stream, width=Inches(2.5))

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
